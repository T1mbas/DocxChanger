# docx_utils.py
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def change_font(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)


def change_spacing(doc):
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def process_document(file_path):
    doc = Document(file_path)
    change_font(doc)
    change_spacing(doc)
    doc.save(file_path)
